"""
Made by Eric Carvalho, https://github.com/ericsiqueira99
18/08/2020
Creates simple essay from wikipedia content
"""


from docx import Document
from docx.shared import Inches
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import messagebox
from urllib.request import urlopen
from bs4 import BeautifulSoup, Tag
import re
from docx2pdf import convert

# GUI
master = tk.Tk()

# get screen width and height
ws = master.winfo_screenwidth()  # width of the screen
hs = master.winfo_screenheight()  # height of the screen

# define app widht and height
w = 320
h = 70

# calculate x and y coordinates for the Tk root window
x = (ws / 2) - (w / 2)
y = (hs / 2) - (h / 2)

# define app settings
app_title = "Essay Maker"
app_icon = "book.ico"
app_size = "%dx%d+%d+%d" % (w, h, x, y)
master.title(app_title)
master.iconbitmap(app_icon)
master.geometry(app_size)

tk.Label(master, text="Choose a topic:").grid(row=0)

e1 = tk.Entry(master)
e1.grid(row=0, column=1)

tk.Button(master,
          text='Make my essay',
          command=lambda: write_essay(e1.get())).grid(row=0, column=3, sticky=tk.W, pady=4)

tk.Button(master, text="help", command=lambda: help_pop_up()).grid(row=1)

save_pdf = tk.IntVar()
tk.Checkbutton(master, text="Save as PDF", variable=save_pdf).grid(row=1, column=1, sticky=tk.W)


# delete file in folder
def delete_file(file):
    if os.path.exists(file):
        os.remove(file)


# converts file from docx to pdf
def save_as_pdf(file):
    convert(file)
    convert(file, file.replace("docx", "pdf"))
    delete_file(file)


# creates pop up window
def help_pop_up():
    messagebox.showinfo("Help", "Enter the topic of your essay and a word document will be automatically"
                                "generated on the same folder as the application."
                                "\nMake sure to enter a valid topic (one that can be retrived from wikipedia)!")


# verifies if <p> content has only bold tags (to avoid importing )
def has_bold_content(tag):
    for child in tag.children:
        if isinstance(child, Tag):
            if child.name == 'b':
                return True
    return False


# controls topic research, cleans text (<p> only) and writes onto word document
def write_essay(topic):
    try:
        source = urlopen('https://en.wikipedia.org/wiki/' + topic).read()
        if topic != "":
            # gather html code from url
            soup = BeautifulSoup(source, 'lxml')
            document = Document()
            title = document.add_heading(topic.title(), 0)
            title.aligment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in soup.find_all('p'):
                if not has_bold_content(paragraph):
                    text = paragraph.text
                    text = re.sub(r'\[.*?\]+', '', text)
                    text = text.replace('\n', '')
                    if len(text.split()) > 30:
                        document.add_paragraph(text)
            doc_name = topic + ".docx"
            document.save(doc_name)
            if save_pdf.get() == 1:
                save_as_pdf(doc_name)
            master.withdraw()
            messagebox.showinfo("All Done :)", "Enjoy your essay")
            master.destroy()
        else:
            messagebox.showerror("Error", "Enter a topic!")
    except:
        e1.delete(first=0, last=30)
        messagebox.showerror("Error", "Invalid topic!")


master.mainloop()
